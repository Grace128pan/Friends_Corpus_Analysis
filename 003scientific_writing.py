import csv
from docx import Document
from mtranslate import translate
import langid

def read_docx(file_path):
    doc = Document(file_path)
    content = []
    for paragraph in doc.paragraphs:
        content.append(paragraph.text)
    return '\n'.join(content)

def translate_text(text, target_lang='en'):
    # Identify the language of the text
    source_lang, _ = langid.classify(text)

    # Skip translation if the source language is already the target language
    if source_lang == target_lang:
        return text

    # Split the text into chunks of 500 characters and translate each chunk
    chunks = [text[i:i+500] for i in range(0, len(text), 500)]
    translated_chunks = [translate(chunk, target_lang) for chunk in chunks]

    # Join the translated chunks to form the complete translation
    translation = ' '.join(translated_chunks)
    return translation

def save_to_docx(original_text, translated_text, output_docx_path):
    doc = Document()
    doc.add_heading('Translation Output', 0)

    doc.add_heading('Original Text (Traditional Chinese)', level=1)
    doc.add_paragraph(original_text)

    doc.add_heading('Translated Text (English)', level=1)
    doc.add_paragraph(translated_text)

    doc.save(output_docx_path)

def main():
    input_file_path = 'C:/Users/grace/Desktop/AQZhengZhuan.docx'
    output_docx_path = 'C:/Users/grace/Desktop/translated_output.docx'

    original_text = read_docx(input_file_path)

    # Translate the text to English using mtranslate
    translated_text = translate_text(original_text, target_lang='en')

    # Save the original and translated text to a Word document
    save_to_docx(original_text, translated_text, output_docx_path)

if __name__ == "__main__":
    main()

