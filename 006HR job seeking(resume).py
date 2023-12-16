import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def format_paragraph(paragraph):
    for run in paragraph.runs:
        run.font.size = Pt(11)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def create_resume(pdf_path, output_path):
    document = Document()

    with fitz.open(pdf_path) as pdf_document:
        for page_number in range(pdf_document.page_count):
            page = pdf_document[page_number]
            text = page.get_text()

            # Add text to the document
            document.add_paragraph(text)

    # Format the document
    for paragraph in document.paragraphs:
        format_paragraph(paragraph)

    # Save the document
    document.save(output_path)

if __name__ == "__main__":
    pdf_path = "C:/Users/grace/Desktop/Resume.pdf"
    output_path = "C:/Users/grace/Desktop/new_resume.docx"

    create_resume(pdf_path, output_path)
